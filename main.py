#!/usr/bin/env/python3.8

import docx
from docx.shared import Pt, Mm
import requests
import parfips
import time


class UserData:
    def __init__(self):
        with open('temp.txt') as file:
            data = file.readlines()
        self.trademarks = data[0].split()
        self.classes = data[1].split()

    def write_docx(self):
        doc = docx.Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        style.paragraph_format.first_line_indent = Mm(10)

        for i in range(len(self.trademarks)):
            tm = parfips.TMData(int(self.trademarks[i]))
            time.sleep(6)
            img = requests.get(tm.get_img_link())

            with open('img.jpg', 'wb') as out:
                out.write(img.content)

            par = doc.add_paragraph()
            run = par.add_run()
            run.add_text(f'{i + 1}. Товарный знак ')
            run.add_picture('img.jpg')
            run.add_text(f' свидетельство № {self.trademarks[i]}, ')
            run.add_text(f'{tm.get_application_date()}, {tm.get_registration_date()}, {tm.get_holder()}, ')
            run.add_text(f'зарегистрирован в отношении {", ".join(tm.get_classes()[1])} классов МКТУ:\n')
            cl = tm.get_classes()
            for j in range(len(cl[0])):
                if cl[1][j] in self.classes:
                    doc.add_paragraph().add_run(f'{cl[0][j]}.')
            if tm.unprotected():
                p = doc.add_paragraph()
                p.add_run('Примечание! ').bold = True
                p.add_run().add_text(f'Неохраняемые элементы товарного знака: {tm.unprotected()}')
            time.sleep(6)

        doc.save('temp.docx')


if __name__ == '__main__':
    u = UserData()
    u.write_docx()
